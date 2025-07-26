"""
Content ingestion pipeline for Universal Knowledge Platform.
Supports multiple file formats with processing and validation.
"""

import asyncio
import logging
import hashlib
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
import mimetypes
import tempfile
import shutil

# File processing libraries
try:
    import PyPDF2
    import docx
    from bs4 import BeautifulSoup
    import requests
except ImportError:
    logging.warning("Some file processing libraries not installed. Install with: pip install PyPDF2 python-docx beautifulsoup4 requests")

from core.config import config
from core.vector_client import vector_db
from core.elasticsearch_client import elasticsearch_client

logger = logging.getLogger(__name__)


@dataclass
class DocumentMetadata:
    """Metadata for ingested documents."""
    title: str
    source: str
    category: str
    author: Optional[str] = None
    created_date: Optional[datetime] = None
    modified_date: Optional[datetime] = None
    language: str = "en"
    tags: List[str] = None
    confidence: float = 1.0
    
    def __post_init__(self):
        if self.tags is None:
            self.tags = []


@dataclass
class IngestedDocument:
    """Represents an ingested document with content and metadata."""
    id: str
    content: str
    metadata: DocumentMetadata
    file_path: Optional[str] = None
    file_size: Optional[int] = None
    mime_type: Optional[str] = None
    ingestion_date: datetime = None
    
    def __post_init__(self):
        if self.ingestion_date is None:
            self.ingestion_date = datetime.utcnow()


class ContentIngestionPipeline:
    """
    Main content ingestion pipeline for processing various file formats.
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.txt': self._process_text,
            '.html': self._process_html,
            '.htm': self._process_html,
            '.md': self._process_markdown,
        }
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.chunk_size = 1000  # characters per chunk
    
    async def ingest_file(
        self,
        file_path: Union[str, Path],
        metadata: Optional[DocumentMetadata] = None
    ) -> IngestedDocument:
        """
        Ingest a single file and return the processed document.
        
        Args:
            file_path: Path to the file to ingest
            metadata: Optional metadata for the document
            
        Returns:
            IngestedDocument with processed content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Validate file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            raise ValueError(f"File too large: {file_size} bytes (max: {self.max_file_size})")
        
        # Get MIME type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        # Generate document ID
        doc_id = self._generate_document_id(file_path, file_size)
        
        # Process file based on extension
        extension = file_path.suffix.lower()
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")
        
        try:
            content = await self.supported_formats[extension](file_path)
            
            # Create metadata if not provided
            if metadata is None:
                metadata = DocumentMetadata(
                    title=file_path.stem,
                    source="file_upload",
                    category="general"
                )
            
            # Create ingested document
            document = IngestedDocument(
                id=doc_id,
                content=content,
                metadata=metadata,
                file_path=str(file_path),
                file_size=file_size,
                mime_type=mime_type
            )
            
            logger.info(f"Successfully ingested file: {file_path}")
            return document
            
        except Exception as e:
            logger.error(f"Failed to ingest file {file_path}: {e}")
            raise
    
    async def ingest_directory(
        self,
        directory_path: Union[str, Path],
        recursive: bool = True,
        file_pattern: str = "*"
    ) -> List[IngestedDocument]:
        """
        Ingest all supported files in a directory.
        
        Args:
            directory_path: Path to directory
            recursive: Whether to process subdirectories
            file_pattern: File pattern to match
            
        Returns:
            List of ingested documents
        """
        directory_path = Path(directory_path)
        documents = []
        
        if recursive:
            files = directory_path.rglob(file_pattern)
        else:
            files = directory_path.glob(file_pattern)
        
        for file_path in files:
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    document = await self.ingest_file(file_path)
                    documents.append(document)
                except Exception as e:
                    logger.error(f"Failed to ingest {file_path}: {e}")
                    continue
        
        logger.info(f"Ingested {len(documents)} documents from {directory_path}")
        return documents
    
    async def ingest_url(
        self,
        url: str,
        metadata: Optional[DocumentMetadata] = None
    ) -> IngestedDocument:
        """
        Ingest content from a URL.
        
        Args:
            url: URL to fetch content from
            metadata: Optional metadata for the document
            
        Returns:
            IngestedDocument with processed content
        """
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            content_type = response.headers.get('content-type', '').lower()
            
            if 'text/html' in content_type:
                content = await self._process_html_content(response.text, url)
            elif 'text/plain' in content_type:
                content = response.text
            else:
                raise ValueError(f"Unsupported content type: {content_type}")
            
            # Create metadata if not provided
            if metadata is None:
                metadata = DocumentMetadata(
                    title=url,
                    source="web_crawl",
                    category="web_content"
                )
            
            # Generate document ID
            doc_id = self._generate_document_id(url, len(content))
            
            document = IngestedDocument(
                id=doc_id,
                content=content,
                metadata=metadata,
                mime_type=content_type
            )
            
            logger.info(f"Successfully ingested URL: {url}")
            return document
            
        except Exception as e:
            logger.error(f"Failed to ingest URL {url}: {e}")
            raise
    
    async def store_document(self, document: IngestedDocument) -> bool:
        """
        Store the ingested document in vector database and search index.
        
        Args:
            document: IngestedDocument to store
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Prepare document for vector storage
            vector_doc = {
                "id": document.id,
                "content": document.content,
                "metadata": {
                    "title": document.metadata.title,
                    "source": document.metadata.source,
                    "category": document.metadata.category,
                    "author": document.metadata.author,
                    "tags": document.metadata.tags,
                    "ingestion_date": document.ingestion_date.isoformat(),
                    "file_path": document.file_path,
                    "mime_type": document.mime_type,
                }
            }
            
            # Store in vector database
            vector_success = await vector_db.store_documents([vector_doc])
            
            # Store in Elasticsearch
            es_success = await elasticsearch_client.index_document(
                doc_id=document.id,
                content=document.content,
                title=document.metadata.title,
                source=document.metadata.source,
                category=document.metadata.category,
                metadata={
                    "author": document.metadata.author,
                    "tags": document.metadata.tags,
                    "ingestion_date": document.ingestion_date.isoformat(),
                    "file_path": document.file_path,
                    "mime_type": document.mime_type,
                }
            )
            
            if vector_success and es_success:
                logger.info(f"Successfully stored document: {document.id}")
                return True
            else:
                logger.error(f"Failed to store document: {document.id}")
                return False
                
        except Exception as e:
            logger.error(f"Error storing document {document.id}: {e}")
            return False
    
    async def process_batch(
        self,
        documents: List[IngestedDocument],
        batch_size: int = 10
    ) -> Dict[str, Any]:
        """
        Process a batch of documents for storage.
        
        Args:
            documents: List of documents to process
            batch_size: Number of documents to process in parallel
            
        Returns:
            Dictionary with processing results
        """
        results = {
            "total": len(documents),
            "successful": 0,
            "failed": 0,
            "errors": []
        }
        
        # Process in batches
        for i in range(0, len(documents), batch_size):
            batch = documents[i:i + batch_size]
            
            # Process batch in parallel
            tasks = [self.store_document(doc) for doc in batch]
            batch_results = await asyncio.gather(*tasks, return_exceptions=True)
            
            for j, result in enumerate(batch_results):
                if isinstance(result, Exception):
                    results["failed"] += 1
                    results["errors"].append({
                        "document_id": batch[j].id,
                        "error": str(result)
                    })
                elif result:
                    results["successful"] += 1
                else:
                    results["failed"] += 1
        
        logger.info(f"Batch processing completed: {results['successful']} successful, {results['failed']} failed")
        return results
    
    def _generate_document_id(self, source: Union[str, Path], size: int) -> str:
        """Generate a unique document ID."""
        content = f"{source}_{size}_{datetime.utcnow().isoformat()}"
        return hashlib.sha256(content.encode()).hexdigest()[:16]
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to process PDF {file_path}: {e}")
    
    async def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to process DOCX {file_path}: {e}")
    
    async def _process_text(self, file_path: Path) -> str:
        """Read text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            raise ValueError(f"Failed to process text file {file_path}: {e}")
    
    async def _process_html(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return await self._process_html_content(file.read(), str(file_path))
        except Exception as e:
            raise ValueError(f"Failed to process HTML file {file_path}: {e}")
    
    async def _process_html_content(self, html_content: str, source: str) -> str:
        """Extract text from HTML content."""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            raise ValueError(f"Failed to process HTML content from {source}: {e}")
    
    async def _process_markdown(self, file_path: Path) -> str:
        """Process markdown file (basic text extraction)."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                # Basic markdown to text conversion
                # Remove markdown syntax
                import re
                content = re.sub(r'#+\s*', '', content)  # Remove headers
                content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Remove bold
                content = re.sub(r'\*(.*?)\*', r'\1', content)  # Remove italic
                content = re.sub(r'`(.*?)`', r'\1', content)  # Remove code
                content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)  # Remove links
                return content.strip()
        except Exception as e:
            raise ValueError(f"Failed to process markdown file {file_path}: {e}")


# CLI interface for testing
async def main():
    """CLI interface for content ingestion."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Content ingestion pipeline")
    parser.add_argument("source", help="File, directory, or URL to ingest")
    parser.add_argument("--title", help="Document title")
    parser.add_argument("--category", help="Document category")
    parser.add_argument("--source", help="Document source")
    parser.add_argument("--tags", nargs="+", help="Document tags")
    parser.add_argument("--recursive", action="store_true", help="Process directories recursively")
    
    args = parser.parse_args()
    
    pipeline = ContentIngestionPipeline()
    
    try:
        if args.source.startswith(('http://', 'https://')):
            # Ingest URL
            metadata = DocumentMetadata(
                title=args.title or args.source,
                source=args.source or "web_crawl",
                category=args.category or "web_content",
                tags=args.tags or []
            )
            document = await pipeline.ingest_url(args.source, metadata)
        elif Path(args.source).is_file():
            # Ingest single file
            metadata = DocumentMetadata(
                title=args.title or Path(args.source).stem,
                source=args.source or "file_upload",
                category=args.category or "general",
                tags=args.tags or []
            )
            document = await pipeline.ingest_file(args.source, metadata)
        elif Path(args.source).is_dir():
            # Ingest directory
            documents = await pipeline.ingest_directory(args.source, args.recursive)
            print(f"Ingested {len(documents)} documents")
            return
        else:
            print(f"Source not found: {args.source}")
            return
        
        # Store document
        success = await pipeline.store_document(document)
        if success:
            print(f"Successfully ingested: {document.id}")
        else:
            print("Failed to store document")
            
    except Exception as e:
        print(f"Error: {e}")


if __name__ == "__main__":
    asyncio.run(main()) 